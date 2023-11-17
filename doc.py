from docx import Document
from concurrent.futures import ThreadPoolExecutor
import os
import fitz

from multiprocessing import Process
class Doc(dict):
    def __init__(self, path) -> None:
        self.name = os.path.basename(path)
        self.path = path
        self.score = {}
        self.text = self.path # tuple with [full_text, headers, bold_underline]
    
    def __str__(self) -> str:
        return f"{self.path}"

    def add_domain(self, key, val):
        self.score[str(key)] = val

    @property
    def name(self):
        return self._name
    @name.setter
    def name(self, name):
        self._name = name
    
    @property
    def path(self):
        return self._path
    @path.setter
    def path(self, path):
        self._path = path

    @property
    def score(self):
        return self._score
    @score.setter
    def score(self, score):
        self._score = score

    @property
    def text(self):
        return self._text
    @text.setter
    def text(self, full_text):
        with ThreadPoolExecutor() as executor:
            full = executor.submit(self.parse_full, full_text)
            bold = executor.submit(self.parse_bold, full_text)
            heading = executor.submit(self.parse_heading, full_text)
            self._text = (full.result(), heading.result(), bold.result()) # tuple of (full text, headers, bold or underline text)

    def parse_full(self, text):
        full_text = ""
        if text.endswith(".docx"):
            document = Document(str(text))
            for para in document.paragraphs:
                full_text += para.text.lower()
        elif text.endswith(".pdf"):
            with fitz.open(str(text)) as pdf:
                for page in pdf: 
                    full_text += page.get_text()
        return full_text

    def parse_bold(self, text):
        bold_text = ""
        if text.endswith(".docx"):
            document = Document(str(text))
            for para in document.paragraphs:
                for run in para.runs:
                    if run.bold or run.underline or run.italic: 
                        bold_text += run.text
        elif text.endswith(".pdf"):
            with fitz.open(str(text)) as pdf: 
                for page in pdf:
                    blocks = page.get_text("dict", flags=11)["blocks"]
                    for b in blocks:  # iterate through the text blocks
                        for l in b["lines"]:  # iterate through the text lines
                            for s in l["spans"]:  # iterate through the text spans
                                if "bold" in s["font"].lower():
                                    bold_text += s["text"]
        return bold_text

    def parse_heading(self, text):
        heading_text = ""
        if text.endswith(".docx"):
            document = Document(str(text))
            for para in document.paragraphs:
                if "Heading" in para.style.name:
                    heading_text += para.text
        return heading_text